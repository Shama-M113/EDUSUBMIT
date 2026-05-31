from flask import Flask, request, render_template, redirect, session, flash, send_from_directory, send_file, jsonify
import os
import sqlite3
import shutil
from datetime import datetime
import docx
import PyPDF2
import io

from sklearn.feature_extraction.text import TfidfVectorizer
from sklearn.metrics.pairwise import cosine_similarity

app = Flask(__name__)
import os

def initialize_database():
    if not os.path.exists('database.db'):
        os.system('python load_student.py')
        os.system('python load_teachers.py')
        os.system('python load_assignments.py')
        os.system('python load_submissions.py')

initialize_database()
app.secret_key = os.environ.get('SECRET_KEY', 'dev-key-change-in-production')
def get_db_connection():
    conn = sqlite3.connect('database.db')
    conn.row_factory = sqlite3.Row  # Enable dictionary-like access to rows
    return conn

@app.route('/')
def home():
    return render_template('home.html')  # Home page with student/teacher buttons

# Student and Teacher Login Pages
@app.route('/student_login')
def student_login_page():
    return render_template('login.html', role='student')

@app.route('/teacher_login')
def teacher_login_page():
    return render_template('login.html', role='teacher')

# Updated Student Login Logic
@app.route('/student_login', methods=['POST'])
def student_login():
    email = request.form.get('email')
    password = request.form.get('password')

    if not email or not password:
        flash("Both email and password fields are required.", "error")
        return redirect('/student_login')

    conn = get_db_connection()
    cursor = conn.cursor()

    # Check if email exists
    cursor.execute("SELECT * FROM students WHERE email = ?", (email,))
    student = cursor.fetchone()

    if student:
        # Check if password matches
        if student['password'] == password:
            session['user_id'] = student['id']
            session['user_name'] = student['name']
            session['user_usn'] = student['usn']
            session['user_email'] = student['email']
            session['role'] = 'student'
            conn.close()
            return redirect('/student_dashboard')
        else:
            flash("Incorrect password.", "error")
    else:
        flash("Email not found.", "error")

    conn.close()
    return redirect('/student_login')

# Updated Teacher Login Logic
@app.route('/teacher_login', methods=['POST'])
def teacher_login():
    email = request.form.get('email')
    password = request.form.get('password')

    if not email or not password:
        flash("Both email and password fields are required.", "error")
        return redirect('/teacher_login')

    conn = get_db_connection()
    cursor = conn.cursor()

    # Check if email exists
    cursor.execute("SELECT * FROM teachers WHERE email = ?", (email,))
    teacher = cursor.fetchone()

    if teacher:
        # Check if password matches
        if teacher['password'] == password:
            session['user_id'] = teacher['teacher_id']
            session['user_name'] = teacher['name']
            session['user_email'] = teacher['email']
            session['role'] = 'teacher'
            conn.close()
            return redirect('/teacher_dashboard')
        else:
            flash("Incorrect password.", "error")
    else:
        flash("Email not found.", "error")

    conn.close()
    return redirect('/teacher_login')

# Student Dashboard
@app.route('/student_dashboard')
def student_dashboard():
    if 'role' not in session or session['role'] != 'student':
        flash("Access denied.", "error")
        return redirect('/student_login')

    conn = get_db_connection()
    cursor = conn.cursor()
    cursor.execute("SELECT id, title, description, deadline FROM assignments")
    assignments = cursor.fetchall()

    # Fetch submitted assignment IDs for this student
    cursor.execute("SELECT assignment_id FROM submissions WHERE student_usn = ?", (session['user_usn'],))
    submitted_ids = [row['assignment_id'] for row in cursor.fetchall()]
    conn.close()

    # Mark each assignment with submission status
    assignment_list = []
    from datetime import datetime
    for a in assignments:
        is_submitted = a['id'] in submitted_ids
        is_past_deadline = datetime.strptime(a['deadline'], '%Y-%m-%d') < datetime.now()

        assignment_list.append({
            **dict(a),
            'is_submitted': is_submitted,
            'is_past_deadline': is_past_deadline
        })

    return render_template('student_dashboard.html',
        name=session['user_name'],
        usn=session['user_usn'],
        email=session['user_email'],
        assignments=assignment_list
    )


# Teacher Dashboard
@app.route('/teacher_dashboard')
def teacher_dashboard():
    if 'role' not in session or session['role'] != 'teacher':
        flash("Access denied. Please log in.", "error")
        return redirect('/teacher_login')

    teacher_name = session['user_name']
    teacher_email = session['user_email']
    tid=session['user_id']

    conn = get_db_connection()
    cursor = conn.cursor()
    cursor.execute("SELECT id, title, description, deadline FROM assignments WHERE teacher_id = ?", (tid,))

    assignments = cursor.fetchall()
    conn.close()

    return render_template('teacher_dashboard.html', name=teacher_name, email=teacher_email, assignments=assignments)

# Create Assignment (Only Teachers)
@app.route('/create_assignment', methods=['GET', 'POST'])
def create_assignment():
    if 'role' not in session or session['role'] != 'teacher':
        flash("Access denied. Please log in as a teacher.", "error")
        return redirect('/teacher_login')

    if request.method == 'POST':
        title = request.form['title']
        description = request.form['description']
        deadline = request.form['deadline']

        if not title or not description or not deadline:
            flash("All fields are required.", "error")
            return redirect('/create_assignment')

        conn = get_db_connection()
        cursor = conn.cursor()
        cursor.execute(
    "INSERT INTO assignments (title, description, deadline, teacher_id) VALUES (?, ?, ?, ?)",
    (title, description, deadline, session['user_id'])
)

        conn.commit()
        conn.close()

        flash(f"Assignment '{title}' created successfully!", "success")
        return redirect('/teacher_dashboard')

    return render_template('create_assignment.html')



from werkzeug.utils import secure_filename
# Student submits assignment

@app.route('/submit_assignment', methods=['POST'])
def submit_assignment():
    assignment_id = request.form.get('assignment_id')
    student_usn = session.get('user_usn')

    if 'file' not in request.files:
        return jsonify({'success': False, 'error': 'No file part'})

    file = request.files['file']
    if file.filename == '':
        return jsonify({'success': False, 'error': 'No selected file'})

    conn = get_db_connection()
    cursor = conn.cursor()

    # Check deadline
    cursor.execute("SELECT deadline FROM assignments WHERE id = ?", (assignment_id,))
    result = cursor.fetchone()
    if result and result["deadline"] < datetime.now().strftime("%Y-%m-%d %H:%M:%S"):
        conn.close()
        return jsonify({'success': False, 'error': 'Deadline passed'})

    # Check if already submitted
    cursor.execute(
        "SELECT * FROM submissions WHERE student_usn = ? AND assignment_id = ?",
        (student_usn, assignment_id)
    )
    if cursor.fetchone():
        conn.close()
        return jsonify({'success': False, 'error': 'Already submitted'})

    # Read file content and store in database
    filename = secure_filename(file.filename)
    file_content = file.read()

    cursor.execute('''
        INSERT INTO submissions (student_usn, student_name, assignment_id, submission_filename, file_content, submission_date)
        VALUES (?, ?, ?, ?, ?, datetime('now'))
    ''', (student_usn, session['user_name'], assignment_id, filename, file_content))
    conn.commit()
    conn.close()

    return jsonify({'success': True})


#view submissions
@app.route('/view_submissions/<int:assignment_id>')
def view_submissions(assignment_id):
    if 'role' not in session or session['role'] != 'teacher':
        flash("Access denied.", "error")
        return redirect('/')
    
    conn = get_db_connection()
    cursor = conn.cursor()
    cursor.execute("SELECT id, assignment_id,student_name,student_usn, submission_date,submission_filename FROM submissions WHERE assignment_id = ?", (assignment_id,))
    submissions = cursor.fetchall()
    conn.close()
    
    return render_template('view_submissions.html', submissions=submissions, assignment_id=assignment_id)
#view file

@app.route('/view_file/<int:submission_id>')
def view_file(submission_id):
    conn = get_db_connection()
    cursor = conn.cursor()

    cursor.execute("SELECT file_content, submission_filename FROM submissions WHERE id = ?", (submission_id,))
    row = cursor.fetchone()
    conn.close()

    if not row or not row['file_content']:
        return "File not found", 404

    return send_file(
        io.BytesIO(row['file_content']),
        as_attachment=True,
        download_name=row['submission_filename']
    )




# ------------------ PLAGIARISM CHECK ------------------

def extract_text(file_path):

    ext = os.path.splitext(file_path)[1].lower()

    # TXT
    if ext == '.txt':
        try:
            with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
                return f.read()
        except:
            return ""

    # DOCX
    elif ext == '.docx':
        try:
            doc = docx.Document(file_path)
            return "\n".join([para.text for para in doc.paragraphs])
        except:
            return ""

    # PDF
    elif ext == '.pdf':
        try:
            text = ""

            with open(file_path, 'rb') as f:
                reader = PyPDF2.PdfReader(f)

                for page in reader.pages:
                    page_text = page.extract_text()

                    if page_text:
                        text += page_text + "\n"

            return text
        except:
            return ""

    return ""


def extract_text_from_bytes(file_content, filename):
    ext = os.path.splitext(filename)[1].lower()

    # Text files (TXT, PY, JS, etc.)
    if ext in ['.txt', '.py', '.js', '.java', '.cpp', '.c', '.html', '.css']:
        try:
            return file_content.decode('utf-8', errors='ignore')
        except:
            return ""

    # DOCX
    elif ext == '.docx':
        try:
            doc = docx.Document(io.BytesIO(file_content))
            return "\n".join([para.text for para in doc.paragraphs])
        except:
            return ""

    # PDF
    elif ext == '.pdf':
        try:
            text = ""
            reader = PyPDF2.PdfReader(io.BytesIO(file_content))
            for page in reader.pages:
                page_text = page.extract_text()
                if page_text:
                    text += page_text + "\n"
            return text
        except:
            return ""

    return ""


@app.route('/check_plagiarism/<int:submission_id>')
def check_plagiarism(submission_id):

    if 'role' not in session or session['role'] != 'teacher':
        flash("Access denied.", "error")
        return redirect('/')

    conn = get_db_connection()
    cursor = conn.cursor()

    # Fetch submission from DB using ID
    cursor.execute(
        "SELECT id, assignment_id, submission_filename, file_content FROM submissions WHERE id = ?",
        (submission_id,)
    )

    row = cursor.fetchone()
    if not row:
        conn.close()
        flash("File not found.", "error")
        return redirect(request.referrer)

    assignment_id = row['assignment_id']
    filename = row['submission_filename']
    file_content = row['file_content']

    # Extract text from file content
    uploaded_text = extract_text_from_bytes(file_content, filename)

    if not uploaded_text.strip():
        conn.close()
        flash("Could not extract readable text.", "error")
        return redirect(request.referrer)

    # Get all other submissions for the SAME assignment
    cursor.execute(
        "SELECT id, submission_filename, file_content FROM submissions WHERE assignment_id = ? AND id != ?",
        (assignment_id, submission_id)
    )
    other_submissions = cursor.fetchall()
    conn.close()

    reference_docs = []
    matched_files = []

    for submission in other_submissions:
        ref_content = submission['file_content']
        ref_filename = submission['submission_filename']

        extracted_text = extract_text_from_bytes(ref_content, ref_filename)

        if extracted_text.strip():
            reference_docs.append(extracted_text)
            matched_files.append(ref_filename)

    if not reference_docs:

        plagiarism_result = "No other submissions available for comparison."

        return render_template(
            'plagiarism_result.html',
            assignment_id=assignment_id,
            filename=filename,
            plagiarism_result=plagiarism_result
        )

    documents = [uploaded_text] + reference_docs

    vectorizer = TfidfVectorizer(stop_words='english')
    tfidf_matrix = vectorizer.fit_transform(documents)

    similarities = cosine_similarity(
        tfidf_matrix[0:1],
        tfidf_matrix[1:]
    )[0]

    max_index = similarities.argmax()
    max_similarity = similarities[max_index] * 100
    matched_file = matched_files[max_index]

    plagiarism_result = (
        f"Plagiarism Score: {round(max_similarity, 2)}% "
        f"(Matched with: {matched_file})"
    )

    if max_similarity > 80:
        plagiarism_result += " - Severe plagiarism detected!"
    elif max_similarity > 50:
        plagiarism_result += " - Moderate plagiarism detected!"
    elif max_similarity > 25:
        plagiarism_result += " - Low similarity detected."
    else:
        plagiarism_result += " - Original content."

    return render_template(
        'plagiarism_result.html',
        assignment_id=assignment_id,
        filename=filename,
        plagiarism_result=plagiarism_result
    )
# ------------------ END ------------------
# Logout
@app.route('/logout')
def logout():
    session.clear()
    return redirect('/')

if __name__ == '__main__':
    app.run(debug=False, host='0.0.0.0', port=int(os.environ.get('PORT', 5000)))
